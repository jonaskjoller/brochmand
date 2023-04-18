from docx import Document

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

doc = input("Type document name:")
# doc = "Test2.docx"
new_doc = doc.split(".")[0]+"_nospaces.docx"
document = Document(doc)
# document_base = document


for i in range(0,len(document.paragraphs)):
    # x = i
    try:
        if document.paragraphs[i].text == '':
            delete_paragraph(document.paragraphs[i])
        document.save(new_doc)
    except:
        
        # print(x,document_base.paragraphs[x-2].text)
        break
